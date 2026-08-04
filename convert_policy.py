from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

# Title
title = doc.add_heading('MANIPUR TECHNICAL UNIVERSITY, IMPHAL', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.runs[0]
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('POLICY & GUIDELINES ON INTERNSHIP/SUMMER TRAINING')
run.font.size = Pt(14)
run.font.bold = True

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run('(At the University and/or at External Host Organisations)')
run.font.size = Pt(12)
run.font.italic = True

doc.add_paragraph()

table = None

with open('Internshipt_Policy_Draft.md', 'r') as f:
    content = f.read()

lines = content.split('\n')
for i, raw_line in enumerate(lines):
    line = raw_line.strip()
    if not line or line == '---':
        if line == '---':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('_' * 50)
            run.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
        continue

    if line.startswith('# ') and 'MANIPUR' in line:
        continue
    elif line.startswith('## '):
        heading = doc.add_heading(line[3:], level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=2)
    elif line.startswith('- ') or line.startswith('* '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('| '):
        if '---' in line:
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if not cells:
            continue
        if table is None:
            table = doc.add_table(rows=1, cols=len(cells))
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            for idx, cell_text in enumerate(cells):
                if idx < len(hdr_cells):
                    hdr_cells[idx].text = cell_text
        else:
            row_cells = table.add_row().cells
            for idx, cell_text in enumerate(cells):
                if idx < len(row_cells):
                    row_cells[idx].text = cell_text
    elif re.match(r'^\d+\.\s', line):
        text = re.sub(r'^\d+\.\s', '', line)
        doc.add_paragraph(text, style='List Number')
    elif line.startswith('```'):
        continue
    elif line:
        p = doc.add_paragraph()
        parts = re.split(r'\*\*|\*', line)
        for idx, part in enumerate(parts):
            if part:
                run = p.add_run(part)
                if idx % 2 == 1:
                    run.font.bold = True

doc.save('Internship_Policy_Draft.docx')
print('Saved Internship_Policy_Draft.docx')
